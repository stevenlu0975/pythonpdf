from typing import List, Dict, Any
import pandas as pd
from docx import Document



def parse_simple_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 1 的方法"""
    # 自定義解析邏輯，返回與 pandaEtlParse 格式一致的結構
    return {"header": data[0], "data": data[1:]}

def parse_education_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 2 的方法"""
    # 使用 DataFrame 做處理
    df = pd.DataFrame(data[1:], columns=data[0])
    return {"header": data[0], "data": df.values.tolist()}

def parse_skill_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 3 的方法"""
    df = pd.DataFrame(data[1:], columns=data[0])
    return {"header": data[0], "data": df.values.tolist()}

def parse_certification_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 4 的方法"""
    columns = data[0]
    records = data[1:]
    df = pd.DataFrame(records, columns=columns)
    df["證書"] = df["證書"].apply(parse_certification_status)
    return {"header": columns, "data": df.values.tolist()}

def parse_complex_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 5 的方法"""
    return {"header": data[0], "data": data[1:]}

def parse_credit_table(data: List[List[Any]]) -> Dict[str, Any]:
    """解析表格 5 的方法"""
    return {"header": data[0], "data": data[1:]}

# 定義處理函數的字典
PARSERS = {
    1: parse_simple_table,
    2: parse_simple_table,
    3: parse_education_table,
    4: parse_skill_table,
    5: parse_certification_table,   
    6: parse_simple_table,
    7: parse_complex_table,
    8: parse_credit_table,
    9: parse_simple_table,
}


# 示例：解析表格
def parse_certification_status(value: str) -> str:
    """示例證書解析邏輯"""
    return "有" if "■有" in value else "無"


def pandaEtlParse(data: List[List[Any]], table_index: int) -> Dict[str, Any]:
    """根據表格索引調用對應的解析函數"""
    if table_index in PARSERS:
        return PARSERS[table_index](data)
    else:
        raise ValueError(f"No parser defined for table index {table_index}")


# 讀取 docx 文件並解析表格
doc = Document("C:\\workspace\\pdf\\test.docx")

print("\n表格內容:")
table_idx = 1
while table_idx <= len(doc.tables):
    table = doc.tables[table_idx - 1]
    print(f"表格 {table_idx}:")
#todo  只有word 應該跑這裡    
    if table_idx == 1:
        data1 = []
        data2 = []
        for row in table.rows:
            data1.append([cell.text for cell in row.cells[:6]])
            data2.append([cell.text for cell in row.cells[7:]])
        
        print(pandaEtlParse(data1, table_idx))
        print(f"表格 {table_idx+1}:")
        print(pandaEtlParse(data2, table_idx+1))
        table_idx += 1
        continue
 
    table_idx += 1
        
 
    data = []
    for row in table.rows:
        data.append([cell.text.strip() for cell in row.cells])
    try:
        parsed_data = pandaEtlParse(data, table_idx)
        print(parsed_data)
    except ValueError as e:
        print(f"Error parsing table {table_idx}: {e}")
